from collections import Counter
from datetime import timedelta
from io import BytesIO

from django.db import IntegrityError
from django.db.models import Avg
from django.utils import timezone
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.http import HttpResponse, JsonResponse
from django.urls import reverse
from django.views import View
from django.views.generic import TemplateView, ListView

from courses.models import Course
from jobs.models import JobAdvert
from jobs.ingestion import extract_advert_metadata, extract_advert_sections, parse_advert_date
from analysis.models import AnalysisRun, GapResult, SkillMatrix, TaskRecord
# Plain functions now — no .delay(), no Celery
from analysis.tasks import (
    run_gap_analysis_task,
    import_csv_task,
    fetch_adzuna_task,
    start_continuous_adzuna_task,
    start_continuous_job_task,
    stop_task,
)


STALE_TASK_MINUTES = 10
BUSINESS_TECHNICAL_SKILL_EXCLUSIONS = {
    "c++",
    "c#",
    "f#",
    "rust",
    "haskell",
    "go",
    "scala",
    "kotlin",
    "java",
    "javascript",
    "typescript",
    "html",
    "css",
    "django",
    "flask",
    "j2ee",
    "z/os",
    "mainframe",
    "kubernetes",
    "devops",
    "ci cd",
    "ci/cd",
    "sre",
    "site reliability engineering",
    "distributed systems",
    "algorithms",
    "encryption",
    "tokenisation",
    "tokenization",
}
BUSINESS_SKILL_LABELS = {
    "data analysis": "business analytics",
    "data analytics": "business analytics",
    "data science": "business analytics",
    "sql": "data-driven decision making",
    "mysql": "data-driven decision making",
    "postgresql": "data-driven decision making",
    "databases": "data governance",
    "python": "analytics automation",
    "r": "analytics automation",
    "machine learning": "AI strategy",
    "artificial intelligence": "AI strategy",
    "generative ai": "AI strategy",
    "genai": "AI strategy",
    "software engineering": "digital transformation",
    "technical architecture": "digital transformation",
    "enterprise architecture": "digital transformation",
    "systems architecture": "digital transformation",
    "cloud native": "digital transformation",
    "cloud-native": "digital transformation",
    "aws": "digital transformation",
    "azure": "digital transformation",
    "gcp": "digital transformation",
    "google cloud": "digital transformation",
    "security": "digital risk governance",
    "web3": "digital transformation",
    "defi": "fintech strategy",
    "fintech": "fintech strategy",
    "automation": "process automation",
    "technical leadership": "digital leadership",
    "technical strategy": "digital strategy",
    "vendor evaluation": "vendor management",
}


def mark_stale_tasks():
    cutoff = timezone.now() - timedelta(minutes=STALE_TASK_MINUTES)
    stale_ids = list(TaskRecord.objects.filter(
        status__in=["PENDING", "STARTED"],
        updated_at__lt=cutoff,
    ).values_list("id", flat=True)[:50])
    if not stale_ids:
        return 0
    return TaskRecord.objects.filter(id__in=stale_ids).update(
        status="FAILURE",
        progress=0,
        notes="Task stopped updating. The background worker likely stopped or the dev server was restarted. Start it again.",
        finished_at=timezone.now(),
    )


def mark_stale_task_if_needed(task):
    if task.status not in ["PENDING", "STARTED"] or not task.updated_at:
        return task
    cutoff = timezone.now() - timedelta(minutes=STALE_TASK_MINUTES)
    if task.updated_at >= cutoff:
        return task
    TaskRecord.objects.filter(
        id=task.id,
        status__in=["PENDING", "STARTED"],
        updated_at__lt=cutoff,
    ).update(
        status="FAILURE",
        progress=0,
        notes="Task stopped updating. The background worker likely stopped or the dev server was restarted. Start it again.",
        finished_at=timezone.now(),
    )
    task.refresh_from_db()
    return task


def task_debug_hint(notes):
    text = (notes or "").lower()
    if "no courses found" in text:
        return "Create a course, add at least one module with content, then start the live pipeline again."
    if "no job adverts found" in text:
        return "Import jobs or let the live pipeline fetch jobs before queueing analysis."
    if "missing adzuna credentials" in text:
        return "Add ADZUNA_APP_ID and ADZUNA_APP_KEY to your .env file, then restart the server."
    if "adzuna api limit reached" in text or "http 429" in text:
        return "The app will retry automatically. Increase the wait interval if this repeats."
    if "adzuna authentication error" in text:
        return "Check ADZUNA_APP_ID and ADZUNA_APP_KEY in .env, then restart the server."
    if "adzuna request error" in text:
        return "Check the keyword, location, and Adzuna country setting."
    if "adzuna network error" in text or "adzuna server error" in text:
        return "This is usually temporary. The jobs-only loop retries automatically."
    if "no text documents" in text:
        return "Add module and job descriptions with enough text for Word2Vec training."
    return "Open Background Tasks for the full task history and check the latest task notes."


def bounded_int(value, default, minimum, maximum):
    try:
        parsed = int(value or default)
    except (TypeError, ValueError):
        parsed = default
    return max(minimum, min(maximum, parsed))


def top_skill_names(counter, limit=5):
    return [skill for skill, _count in counter.most_common(limit)]


def refine_business_skill(skill):
    normalized = (skill or "").strip().lower()
    if normalized in BUSINESS_TECHNICAL_SKILL_EXCLUSIONS:
        return None
    return BUSINESS_SKILL_LABELS.get(normalized, skill)


def refined_business_counter(counter):
    refined = Counter()
    for skill, count in counter.items():
        label = refine_business_skill(skill)
        if label:
            refined[label] += count
    return refined


def refine_skill_rows_for_business(rows, limit=10):
    refined = Counter()
    for row in rows:
        label = refine_business_skill(row["skill"])
        if label:
            refined[label] += row["frequency"]
    return [
        {"skill": skill, "frequency": frequency}
        for skill, frequency in refined.most_common(limit)
    ]


def join_skill_names(skills):
    if not skills:
        return ""
    if len(skills) == 1:
        return skills[0]
    return f"{', '.join(skills[:-1])} and {skills[-1]}"


def recommendation_skill_insight(missing_counter, matched_counter):
    missing = top_skill_names(refined_business_counter(missing_counter))
    matched = top_skill_names(refined_business_counter(matched_counter), 4)
    parts = []
    if missing:
        parts.append(f"Based on the analysed course-to-job data, the curriculum should strengthen {join_skill_names(missing)}.")
    if matched:
        parts.append(f"The current evidence already shows coverage in {join_skill_names(matched)}.")
    if not parts:
        parts.append("No repeated skill pattern is visible yet, so inspect the lowest-scoring job matches first.")
    return " ".join(parts)


def build_skill_suggestion_matrix(results, limit=8):
    matrix = {}
    for result in results:
        matched_labels = {
            label for label in (refine_business_skill(skill) for skill in (result.matched_skills or []))
            if label
        }
        missing_labels = {
            label for label in (refine_business_skill(skill) for skill in (result.missing_skills or []))
            if label
        }
        for skill in matched_labels | missing_labels:
            item = matrix.setdefault(skill, {
                "job_ids": set(),
                "covered_course_ids": set(),
                "matched_pairs": set(),
                "missing_pairs": set(),
            })
            item["job_ids"].add(result.job_id)
            pair = (result.course_id, result.job_id)
            if skill in matched_labels:
                item["covered_course_ids"].add(result.course_id)
                item["matched_pairs"].add(pair)
            if skill in missing_labels:
                item["missing_pairs"].add(pair)

    max_gap = max((len(item["missing_pairs"]) for item in matrix.values()), default=1)
    rows = []
    for skill, item in matrix.items():
        matched_count = len(item["matched_pairs"])
        missing_count = len(item["missing_pairs"])
        demand_count = matched_count + missing_count
        coverage = round((matched_count / max(1, demand_count)) * 100)
        gap_percent = round((missing_count / max(1, demand_count)) * 100)
        rows.append({
            "skill": skill,
            "demand_count": demand_count,
            "job_count": len(item["job_ids"]),
            "course_count": len(item["covered_course_ids"]),
            "matched_count": matched_count,
            "missing_count": missing_count,
            "coverage": coverage,
            "gap_percent": gap_percent,
            "gap_width": round((missing_count / max_gap) * 100),
        })
    return sorted(
        rows,
        key=lambda row: (-row["missing_count"], row["coverage"], row["skill"]),
    )[:limit]


class DashboardView(TemplateView):
    template_name = "dashboard/home.html"

    def get_context_data(self, **kwargs):
        mark_stale_tasks()
        ctx = super().get_context_data(**kwargs)
        ctx["course_count"] = Course.objects.count()
        ctx["module_count"] = sum(c.modules.count() for c in Course.objects.prefetch_related("modules"))
        ctx["job_count"] = JobAdvert.objects.count()
        ctx["last_run"] = AnalysisRun.objects.first()
        ctx["pending_tasks"] = TaskRecord.objects.filter(status__in=["PENDING", "STARTED"]).count()
        ctx["live_task"] = (
            TaskRecord.objects
            .filter(run_name__startswith="Jobs Only", status__in=["PENDING", "STARTED"])
            .first()
            or TaskRecord.objects
            .filter(run_name__startswith="Live Pipeline", status__in=["PENDING", "STARTED"])
            .first()
        )
        latest_jobs_only = TaskRecord.objects.filter(run_name__startswith="Jobs Only").order_by("-created_at").first()
        ctx["should_autostart_jobs"] = not ctx["live_task"] and not (latest_jobs_only and latest_jobs_only.status == "STOPPED")
        ctx["avg_score"] = (GapResult.objects.aggregate(v=Avg("similarity_score"))["v"] or 0) * 100
        ctx["latest_results"] = GapResult.objects.select_related("course", "job").order_by("-run__created_at", "-similarity_score")[:8]
        ctx["network_schools"] = (
            Course.objects
            .exclude(university_name="")
            .order_by("university_name")
            .values_list("university_name", flat=True)
            .distinct()
        )
        ctx["network_jobs"] = JobAdvert.objects.order_by("title").values("id", "title", "company")[:500]
        return ctx


class JobUploadView(View):
    template_name = "jobs/upload.html"

    def get(self, request):
        return render(request, self.template_name)

    def post(self, request):
        t = request.POST.get("upload_type", "csv")

        if t == "csv":
            f = request.FILES.get("csv_file")
            if not f:
                messages.error(request, "Select a CSV file first.")
                return render(request, self.template_name)
            raw = f.read()          # bytes — passed directly, no list() conversion needed
            record = TaskRecord.objects.create(run_name=f"CSV: {f.name}")
            import_csv_task(raw, record_id=record.id)   # fires thread, returns immediately
            messages.success(request, f"'{f.name}' is being imported in the background — track progress under Tasks.")
            return redirect("task-list")

        elif t == "adzuna":
            kw = request.POST.get("keyword", "").strip()
            loc = request.POST.get("location", "south africa").strip()
            n = int(request.POST.get("max_results", 50))
            if not kw:
                messages.error(request, "Enter a keyword.")
                return render(request, self.template_name)
            record = TaskRecord.objects.create(run_name=f"Adzuna: {kw}")
            fetch_adzuna_task(kw, loc, n, record_id=record.id)
            messages.success(request, f"Fetching '{kw}' jobs in the background.")
            return redirect("task-list")

        elif t == "manual":
            title = request.POST.get("title", "").strip()
            desc = request.POST.get("description", "").strip()
            if not title or not desc:
                messages.error(request, "Title and description are required.")
                return render(request, self.template_name)
            sections = extract_advert_sections(desc)
            metadata = extract_advert_metadata(desc)
            try:
                JobAdvert.objects.create(
                    title=title,
                    company=request.POST.get("company", "").strip(),
                    recruiter=request.POST.get("recruiter", "").strip() or metadata.get("recruiter", ""),
                    job_reference=request.POST.get("job_reference", "").strip() or metadata.get("job_reference", ""),
                    location=request.POST.get("location", "").strip() or metadata.get("location", ""),
                    category=request.POST.get("category", "").strip(),
                    contract_type=request.POST.get("contract_type", "").strip(),
                    contract_time=request.POST.get("contract_time", "").strip(),
                    summary=request.POST.get("summary", "").strip() or sections["summary"],
                    position_info=request.POST.get("position_info", "").strip() or sections["position_info"],
                    description=desc,
                    source="upload",
                    date_posted=parse_advert_date(metadata.get("date_posted")),
                )
            except IntegrityError:
                messages.warning(request, "That job advert already exists, so it was not added again.")
                return redirect("job-list")
            messages.success(request, f"Job '{title}' added.")
            return redirect("job-list")

        return redirect("job-list")


class JobListView(ListView):
    model = JobAdvert
    template_name = "jobs/list.html"
    context_object_name = "jobs"
    paginate_by = 30


class JobDeleteView(View):
    def post(self, request, pk):
        get_object_or_404(JobAdvert, pk=pk).delete()
        messages.success(request, "Job deleted.")
        return redirect("job-list")


class RunAnalysisView(View):
    def post(self, request):
        name = request.POST.get("run_name", "Analysis Run")
        max_jobs = bounded_int(request.POST.get("max_jobs"), 0, 0, 100000) or None
        if max_jobs and "Smoke" not in name:
            name = f"Smoke {name}"
        record = TaskRecord.objects.create(run_name=name)
        run_gap_analysis_task(run_name=name, record_id=record.id, max_jobs=max_jobs)   # fires thread
        if request.headers.get("x-requested-with") == "XMLHttpRequest":
            return JsonResponse({
                "task_id": record.id,
                "status_url": reverse("task-status-api", args=[record.id]),
                "task_url": reverse("task-list"),
                "max_jobs": max_jobs,
            })
        messages.success(request, f"Analysis '{name}' started in the background.")
        return redirect("task-list")


class StartContinuousJobsView(View):
    def post(self, request):
        mark_stale_tasks()
        live = TaskRecord.objects.filter(run_name__startswith="Live Pipeline", status__in=["PENDING", "STARTED"]).first()
        if live:
            return JsonResponse({
                "task_id": live.id,
                "status_url": reverse("task-status-api", args=[live.id]),
            })

        keyword = request.POST.get("keyword", "MBA").strip() or "MBA"
        location = request.POST.get("location", "south africa").strip() or "south africa"
        max_results = bounded_int(request.POST.get("max_results"), 50, 1, 50)
        interval = bounded_int(request.POST.get("interval_seconds"), 45, 10, 3600)
        record = TaskRecord.objects.create(run_name=f"Live Pipeline: {keyword}")
        start_continuous_job_task(keyword, location, max_results, interval, record_id=record.id)
        return JsonResponse({
            "task_id": record.id,
            "status_url": reverse("task-status-api", args=[record.id]),
        })


class StartJobsOnlyView(View):
    def post(self, request):
        mark_stale_tasks()
        live = TaskRecord.objects.filter(run_name__startswith="Jobs Only", status__in=["PENDING", "STARTED"]).first()
        if live:
            return JsonResponse({
                "task_id": live.id,
                "status_url": reverse("task-status-api", args=[live.id]),
            })

        keyword = request.POST.get("keyword", "MBA").strip() or "MBA"
        location = request.POST.get("location", "south africa").strip() or "south africa"
        max_results = bounded_int(request.POST.get("max_results"), 50, 1, 50)
        interval = bounded_int(request.POST.get("interval_seconds"), 45, 5, 3600)
        record = TaskRecord.objects.create(run_name=f"Jobs Only: {keyword}")
        start_continuous_adzuna_task(keyword, location, max_results, interval, record_id=record.id)
        return JsonResponse({
            "task_id": record.id,
            "status_url": reverse("task-status-api", args=[record.id]),
        })


class StopTaskView(View):
    def post(self, request, pk):
        stop_task(pk)
        if request.headers.get("x-requested-with") == "XMLHttpRequest":
            return JsonResponse({"ok": True})
        messages.success(request, "Pause requested.")
        return redirect("home")


class AnalysisResultsView(TemplateView):
    template_name = "analysis/results.html"

    def get_context_data(self, **kwargs):
        ctx = super().get_context_data(**kwargs)
        runs = AnalysisRun.objects.all()
        ctx["runs"] = runs

        run_id = self.request.GET.get("run")
        sel = None
        if run_id:
            sel = get_object_or_404(AnalysisRun, id=run_id)
        elif runs.exists():
            sel = runs.first()

        if sel:
            school_filter = self.request.GET.get("school", "").strip()
            threshold = bounded_int(self.request.GET.get("threshold"), 55, 0, 100)
            all_results_qs = (
                GapResult.objects
                .filter(run=sel)
                .select_related("course", "job")
                .order_by("-similarity_score")
            )
            if school_filter:
                all_results_qs = all_results_qs.filter(course__university_name=school_filter)
            all_results = list(all_results_qs)

            ctx["selected_run"] = sel
            ctx["selected_school"] = school_filter
            ctx["threshold"] = threshold
            ctx["schools"] = (
                Course.objects
                .filter(gapresult__run=sel)
                .exclude(university_name="")
                .order_by("university_name")
                .values_list("university_name", flat=True)
                .distinct()
            )
            ctx["results"] = all_results[:300]
            ctx["job_skills"] = SkillMatrix.objects.filter(run=sel, source="jobs")[:20]
            ctx["course_skills"] = SkillMatrix.objects.filter(run=sel, source="courses")[:20]
            visual_data = build_results_visual_data(all_results, threshold)
            ctx.update(visual_data)
        return ctx


def build_results_visual_data(results, threshold):
    bands = [
        {"key": "0-20", "low": 0, "high": 20},
        {"key": "20-40", "low": 20, "high": 40},
        {"key": "40-60", "low": 40, "high": 60},
        {"key": "60-80", "low": 60, "high": 80},
        {"key": "80-100", "low": 80, "high": 101},
    ]
    course_map = {}
    school_map = {}
    max_cell = 1

    for result in results:
        course = result.course
        school = course.university_name or "Unassigned school"
        score = result.similarity_percent
        course_item = course_map.setdefault(course.id, {
            "course": course,
            "school": school,
            "scores": [],
            "matched_total": 0,
            "missing_total": 0,
            "matched_skills": Counter(),
            "missing_skills": Counter(),
            "cells": {band["key"]: 0 for band in bands},
        })
        course_item["scores"].append(score)
        course_item["matched_total"] += len(result.matched_skills or [])
        course_item["missing_total"] += len(result.missing_skills or [])
        course_item["matched_skills"].update(result.matched_skills or [])
        course_item["missing_skills"].update(result.missing_skills or [])
        for band in bands:
            if band["low"] <= score < band["high"]:
                course_item["cells"][band["key"]] += 1
                max_cell = max(max_cell, course_item["cells"][band["key"]])
                break

        school_item = school_map.setdefault(school, {
            "scores": [],
            "matched_total": 0,
            "missing_total": 0,
            "matched_skills": Counter(),
            "missing_skills": Counter(),
            "courses": set(),
        })
        school_item["scores"].append(score)
        school_item["matched_total"] += len(result.matched_skills or [])
        school_item["missing_total"] += len(result.missing_skills or [])
        school_item["matched_skills"].update(result.matched_skills or [])
        school_item["missing_skills"].update(result.missing_skills or [])
        school_item["courses"].add(course.id)

    heatmap_rows = []
    scatter_points = []
    course_recommendations = []
    for item in sorted(course_map.values(), key=lambda value: value["course"].name):
        scores = item["scores"]
        avg_score = round(sum(scores) / max(1, len(scores)), 1)
        mismatch = avg_score < threshold
        cells = []
        for band in bands:
            count = item["cells"][band["key"]]
            intensity = count / max_cell if max_cell else 0
            cells.append({"band": band["key"], "count": count, "intensity": round(intensity, 3)})
        heatmap_rows.append({
            "course": item["course"],
            "school": item["school"],
            "avg_score": avg_score,
            "mismatch": mismatch,
            "cells": cells,
        })
        scatter_points.append({
            "label": item["course"].code or item["course"].name,
            "course": item["course"].name,
            "school": item["school"],
            "x": item["matched_total"],
            "y": item["missing_total"],
            "score": avg_score,
        })
        if mismatch:
            skill_insight = recommendation_skill_insight(item["missing_skills"], item["matched_skills"])
            suggested_skills = join_skill_names(top_skill_names(refined_business_counter(item["missing_skills"])))
            course_recommendations.append({
                "label": item["course"].code or item["course"].name,
                "school": item["school"],
                "score": avg_score,
                "skills": suggested_skills,
                "message": f"For this programme, the data suggests a curriculum refresh around {suggested_skills or 'the lowest-covered demand skills'}. {skill_insight} Review module outcomes, readings, projects, and assessment language around those demand areas.",
            })

    school_recommendations = []
    school_summaries = []
    for school, item in sorted(school_map.items()):
        scores = item["scores"]
        avg_score = round(sum(scores) / max(1, len(scores)), 1)
        matched = item["matched_total"]
        missing = item["missing_total"]
        mismatch = avg_score < threshold or missing > matched
        summary = {
            "school": school,
            "avg_score": avg_score,
            "matched_total": matched,
            "missing_total": missing,
            "course_count": len(item["courses"]),
            "mismatch": mismatch,
        }
        school_summaries.append(summary)
        if mismatch:
            if avg_score < threshold:
                reason = f"Average score is {avg_score}%, below the {threshold}% threshold."
            else:
                reason = f"Missing skill evidence ({missing}) exceeds matched evidence ({matched})."
            skill_insight = recommendation_skill_insight(item["missing_skills"], item["matched_skills"])
            school_recommendations.append({
                "school": school,
                "score": avg_score,
                "message": f"{reason} {skill_insight} Prioritise curriculum updates in modules linked to the job adverts with the highest missing-skill counts.",
            })

    if not school_recommendations and school_summaries:
        best = max(school_summaries, key=lambda item: item["avg_score"])
        school_recommendations.append({
            "school": best["school"],
            "score": best["avg_score"],
            "message": "No curriculum is below the current threshold. Use the scatter plot to inspect outlier courses with high missing-skill counts before changing module content.",
        })

    return {
        "score_bands": [band["key"] for band in bands],
        "heatmap_rows": heatmap_rows,
        "scatter_points": scatter_points,
        "school_summaries": school_summaries,
        "school_recommendations": school_recommendations,
        "course_recommendations": course_recommendations[:12],
        "skill_suggestion_matrix": build_skill_suggestion_matrix(results),
    }


def docx_add_field(paragraph, instruction, placeholder="Right-click and update field in Word."):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)

    paragraph.add_run(placeholder)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def docx_shade(cell, fill):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def docx_set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold


def add_key_value_table(document, rows):
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    docx_set_cell_text(hdr[0], "Metric", True)
    docx_set_cell_text(hdr[1], "Value", True)
    docx_shade(hdr[0], "F2EEE8")
    docx_shade(hdr[1], "F2EEE8")
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(key)
        cells[1].text = str(value)
    return table


def add_simple_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        docx_set_cell_text(table.rows[0].cells[index], header, True)
        docx_shade(table.rows[0].cells[index], "F2EEE8")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
    return table


def score_fill(score):
    if score >= 80:
        return "F58220"
    if score >= 60:
        return "FFBD80"
    if score >= 40:
        return "FFE1C2"
    if score > 0:
        return "F1E7DD"
    return "F1F3F5"


def add_cross_tab_table(document, results, limit_courses=12, limit_jobs=10):
    courses = []
    jobs = []
    for result in results:
        if result.course not in courses:
            courses.append(result.course)
        if result.job not in jobs:
            jobs.append(result.job)
    courses = courses[:limit_courses]
    jobs = jobs[:limit_jobs]
    if not courses or not jobs:
        document.add_paragraph("No cross-tab data available for this school.")
        return

    score_map = {(result.course_id, result.job_id): result.similarity_percent for result in results}
    headers = ["Course"] + [job.title[:28] for job in jobs]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        docx_set_cell_text(table.rows[0].cells[index], header, True)
        docx_shade(table.rows[0].cells[index], "F2EEE8")
    for course in courses:
        cells = table.add_row().cells
        cells[0].text = course.code or course.name[:32]
        for index, job in enumerate(jobs, start=1):
            score = score_map.get((course.id, job.id), 0)
            cells[index].text = f"{score:.0f}" if score else ""
            docx_shade(cells[index], score_fill(score))


def add_skill_matrix_table(document, matrix_rows):
    if not matrix_rows:
        document.add_paragraph("No school skill matrix data available.")
        return
    rows = [
        [
            row["skill"],
            row["demand_count"],
            row["job_count"],
            row["course_count"],
            row["missing_count"],
            f'{row["gap_percent"]}%',
            f'{row["coverage"]}%',
        ]
        for row in matrix_rows
    ]
    add_simple_table(
        document,
        ["Skill", "Demand evidence", "Jobs", "Courses covered", "Gap evidence", "Gap %", "Coverage %"],
        rows,
    )


def latest_visual_run():
    run = AnalysisRun.objects.first()
    if run and not GapResult.objects.filter(run=run).exists():
        run = AnalysisRun.objects.filter(results__isnull=False).distinct().order_by("-created_at").first()
    return run


class TechnicalReportExportView(View):
    def get(self, request):
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return HttpResponse("python-docx is not installed.", status=500)

        run = latest_visual_run()
        schools = list(
            Course.objects
            .exclude(university_name="")
            .order_by("university_name")
            .values_list("university_name", flat=True)
            .distinct()
        )
        if Course.objects.filter(university_name="").exists():
            schools.append("Unassigned school")

        document = Document()
        section = document.sections[0]
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("Page ")
        docx_add_field(footer, "PAGE", "1")

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("CurriculumMatch Technical Report")
        title_run.bold = True
        title_run.font.size = None
        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.add_run("Curriculum-to-job-market alignment report").italic = True
        document.add_paragraph(f"Generated: {timezone.localtime(timezone.now()).strftime('%d %B %Y %H:%M')}")
        document.add_paragraph(f"Analysis run: {run.name if run else 'No completed analysis run available'}")
        document.add_paragraph("Scope: all schools currently stored in the database.")
        document.add_page_break()

        document.add_heading("Table of Contents", level=1)
        docx_add_field(document.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u')
        document.add_paragraph("Note: in Microsoft Word, right-click the table above and choose Update Field to refresh page numbers.")
        document.add_page_break()

        document.add_heading("1. Introduction", level=1)
        document.add_paragraph(
            "This technical report summarises curriculum-to-job-market alignment using the courses, modules, "
            "job adverts, extracted skills, semantic similarity scores, and school metadata currently stored in CurriculumMatch. "
            "It is designed as an operational report, not a literature review."
        )

        document.add_heading("2. Methodology", level=1)
        document.add_paragraph(
            "The system parses curriculum and job evidence, extracts explicit skill terms, creates semantic vectors, "
            "and compares each course against job adverts. The final score blends semantic similarity with explicit skill coverage."
        )
        document.add_paragraph("Cosine similarity: cos(A, B) = (A . B) / (||A|| * ||B||)")
        document.add_paragraph("Course semantic score: mean(top-k module-to-job cosine scores)")
        document.add_paragraph("Skill coverage: matched job skills / unique job skills")
        document.add_paragraph("Final score: ((0.75 * semantic) + (0.25 * skill coverage)) / (0.75 + 0.25)")
        document.add_paragraph(
            "School Skill Matrix: demand evidence is the count of course-to-job comparisons where a refined skill appears as "
            "matched or missing. Gap percentage is gap evidence divided by demand evidence. Coverage percentage is covered "
            "evidence divided by demand evidence."
        )

        document.add_heading("3. Dashboard Snapshot", level=1)
        all_results = list(GapResult.objects.filter(run=run).select_related("course", "job")) if run else []
        avg_score = round((sum(result.similarity_score for result in all_results) / len(all_results)) * 100, 1) if all_results else 0
        add_key_value_table(document, [
            ("Schools", len(schools)),
            ("Courses", Course.objects.count()),
            ("Job adverts", JobAdvert.objects.count()),
            ("Course-job comparisons", len(all_results)),
            ("Average final score", f"{avg_score}%"),
        ])

        document.add_heading("4. School Results", level=1)
        for school in schools:
            document.add_heading(school, level=2)
            if school == "Unassigned school":
                school_results = [result for result in all_results if not result.course.university_name]
            else:
                school_results = [result for result in all_results if result.course.university_name == school]
            school_courses = Course.objects.filter(university_name="" if school == "Unassigned school" else school)
            visual_data = build_results_visual_data(school_results, 55)
            summary = visual_data["school_summaries"][0] if visual_data["school_summaries"] else None
            add_key_value_table(document, [
                ("Courses in database", school_courses.count()),
                ("Analysed comparisons", len(school_results)),
                ("Average score", f'{summary["avg_score"]}%' if summary else "No analysis data"),
                ("Matched evidence", summary["matched_total"] if summary else 0),
                ("Missing evidence", summary["missing_total"] if summary else 0),
                ("Mismatch risk", "Yes" if summary and summary["mismatch"] else "No"),
            ])

            document.add_heading("Curriculum Recommendations", level=3)
            for item in visual_data["school_recommendations"]:
                document.add_paragraph(f'{item["school"]} ({item["score"]}%): {item["message"]}', style=None)

            document.add_heading("School Skill Matrix", level=3)
            add_skill_matrix_table(document, visual_data["skill_suggestion_matrix"])

            document.add_heading("Course-to-Job Cross-tab Snapshot", level=3)
            add_cross_tab_table(document, school_results)

            document.add_heading("Top Course-to-Job Matches", level=3)
            top_rows = [
                [
                    result.course.name[:48],
                    result.job.title[:48],
                    result.job.company or "",
                    f"{result.similarity_percent}%",
                    ", ".join((result.matched_skills or [])[:5]),
                    ", ".join((result.missing_skills or [])[:5]),
                ]
                for result in sorted(school_results, key=lambda item: item.similarity_score, reverse=True)[:8]
            ]
            if top_rows:
                add_simple_table(document, ["Course", "Job advert", "Company", "Score", "Matched", "Missing"], top_rows)
            else:
                document.add_paragraph("No analysed matches available for this school.")

        document.add_heading("5. Appendix: Interpretation Notes", level=1)
        document.add_paragraph(
            "Scores and evidence counts are decision-support signals. Low alignment can indicate a genuine curriculum gap, "
            "a job advert outside programme scope, thin module text, or terminology differences. Recommendations should be "
            "reviewed by curriculum owners before module changes are made."
        )

        output = BytesIO()
        document.save(output)
        output.seek(0)
        response = HttpResponse(
            output.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = 'attachment; filename="curriculummatch-technical-report.docx"'
        return response


class TaskListView(ListView):
    model = TaskRecord
    template_name = "analysis/tasks.html"
    context_object_name = "tasks"
    paginate_by = 40

    def get_context_data(self, **kwargs):
        ctx = super().get_context_data(**kwargs)
        ctx["has_running"] = TaskRecord.objects.filter(status__in=["PENDING", "STARTED"]).exists()
        return ctx


def task_status_api(request, pk):
    r = get_object_or_404(TaskRecord, pk=pk)
    r = mark_stale_task_if_needed(r)
    return JsonResponse({
        "id": r.id,
        "run_name": r.run_name,
        "status": r.status,
        "progress": r.progress,
        "notes": r.notes,
        "debug_hint": task_debug_hint(r.notes) if r.status == "FAILURE" else "",
        "updated_at": r.updated_at.isoformat() if r.updated_at else None,
        "finished_at": r.finished_at.isoformat() if r.finished_at else None,
    })


def dashboard_metrics(request):
    last_run = AnalysisRun.objects.first()
    visual_run = last_run
    if visual_run and not GapResult.objects.filter(run=visual_run).exists():
        visual_run = AnalysisRun.objects.filter(results__isnull=False).distinct().order_by("-created_at").first()
    results = GapResult.objects.filter(run=visual_run).select_related("course", "job") if visual_run else GapResult.objects.none()
    score_values = list(results.values_list("similarity_score", flat=True))
    buckets = [0, 0, 0, 0, 0]
    for score in score_values:
        idx = min(4, int(max(0, score) * 5))
        buckets[idx] += 1

    job_skills = []
    course_skills = []
    if visual_run:
        job_skill_rows = list(SkillMatrix.objects.filter(run=visual_run, source="jobs").values("skill", "frequency")[:30])
        course_skill_rows = list(SkillMatrix.objects.filter(run=visual_run, source="courses").values("skill", "frequency")[:30])
        job_skills = refine_skill_rows_for_business(job_skill_rows)
        course_skills = refine_skill_rows_for_business(course_skill_rows)

    recent_tasks = []
    for task in TaskRecord.objects.order_by("-created_at").values("id", "run_name", "status", "progress", "notes")[:8]:
        task["debug_hint"] = task_debug_hint(task.get("notes")) if task["status"] == "FAILURE" else ""
        recent_tasks.append(task)
    return JsonResponse({
        "counts": {
            "courses": Course.objects.count(),
            "modules": sum(c.modules.count() for c in Course.objects.prefetch_related("modules")),
            "jobs": JobAdvert.objects.count(),
            "runs": AnalysisRun.objects.count(),
        },
        "last_run": {
            "id": last_run.id,
            "status": last_run.status,
            "name": last_run.name,
        } if last_run else None,
        "visual_run": {
            "id": visual_run.id,
            "status": visual_run.status,
            "name": visual_run.name,
        } if visual_run else None,
        "has_visual_data": bool(score_values or job_skills or course_skills),
        "average_score": round((sum(score_values) / len(score_values)) * 100, 1) if score_values else 0,
        "score_buckets": buckets,
        "score_labels": ["0-20", "20-40", "40-60", "60-80", "80-100"],
        "job_skills": job_skills,
        "course_skills": course_skills,
        "recent_tasks": recent_tasks,
    })


def similarity_network(request):
    try:
        import networkx as nx
    except ImportError:
        nx = None

    last_run = AnalysisRun.objects.first()
    visual_run = last_run
    if visual_run and not GapResult.objects.filter(run=visual_run).exists():
        visual_run = AnalysisRun.objects.filter(results__isnull=False).distinct().order_by("-created_at").first()
    school_filter = request.GET.get("school", "").strip()
    job_filter = bounded_int(request.GET.get("job"), 0, 0, 100000000) or None
    qs = GapResult.objects.none()
    if visual_run:
        qs = GapResult.objects.filter(run=visual_run).select_related("course", "job")
        if school_filter:
            qs = qs.filter(course__university_name=school_filter)
        if job_filter:
            qs = qs.filter(job_id=job_filter)
        qs = qs.order_by("-similarity_score")[:40]
    if nx:
        graph = nx.Graph()
        for r in qs:
            course_id = f"course-{r.course_id}"
            job_id = f"job-{r.job_id}"
            graph.add_node(course_id, label=r.course.code or r.course.name, group="course", title=r.course.name)
            graph.add_node(job_id, label=r.job.title[:28], group="job", title=r.job.title)
            graph.add_edge(course_id, job_id, value=max(1, r.similarity_percent), title=f"Cosine similarity: {r.similarity_score:.4f}")
        positions = nx.spring_layout(graph, seed=42, k=0.7) if graph.number_of_nodes() else {}
        nodes = [
            {
                "id": n,
                "x": round(positions.get(n, (0, 0))[0] * 520, 2),
                "y": round(positions.get(n, (0, 0))[1] * 360, 2),
                **attrs,
            }
            for n, attrs in graph.nodes(data=True)
        ]
        edges = [{"from": u, "to": v, **attrs} for u, v, attrs in graph.edges(data=True)]
    else:
        nodes, edges, seen = [], [], set()
        for index, r in enumerate(qs):
            course_id = f"course-{r.course_id}"
            job_id = f"job-{r.job_id}"
            if course_id not in seen:
                nodes.append({"id": course_id, "label": r.course.code or r.course.name, "group": "course", "title": r.course.name, "x": -260, "y": index * 52})
                seen.add(course_id)
            if job_id not in seen:
                nodes.append({"id": job_id, "label": r.job.title[:28], "group": "job", "title": r.job.title, "x": 260, "y": index * 52})
                seen.add(job_id)
            edges.append({"from": course_id, "to": job_id, "value": max(1, r.similarity_percent), "title": f"Cosine similarity: {r.similarity_score:.4f}"})
    return JsonResponse({
        "run_id": visual_run.id if visual_run else None,
        "has_visual_data": bool(nodes and edges),
        "nodes": nodes,
        "edges": edges,
    })


def results_json(request):
    run_id = request.GET.get("run")
    if not run_id:
        return JsonResponse({"error": "run param required"}, status=400)
    qs = GapResult.objects.filter(run_id=run_id).select_related("course", "job")
    return JsonResponse({"results": [
        {
            "course": r.course.name,
            "job": r.job.title,
            "company": r.job.company,
            "recruiter": r.job.recruiter,
            "job_reference": r.job.job_reference,
            "location": r.job.location,
            "category": r.job.category,
            "contract_type": r.job.contract_type,
            "contract_time": r.job.contract_time,
            "date_posted": r.job.date_posted.isoformat() if r.job.date_posted else None,
            "summary": r.job.summary,
            "position_info": r.job.position_info,
            "score": round(r.similarity_score * 100, 1),
            "matched": r.matched_skills,
            "missing": r.missing_skills,
        } for r in qs
    ]})
